import os
import re
import streamlit as st
import concurrent.futures
import torch


from googlesearch import search
import feedparser

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from sentence_transformers import SentenceTransformer
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import CrossEncoderReranker
from langchain_community.embeddings.huggingface import HuggingFaceBgeEmbeddings
from langchain_community.cross_encoders import HuggingFaceCrossEncoder

from langchain.embeddings.base import Embeddings

from dotenv import load_dotenv

import google.generativeai as genai
from google.generativeai import types
from pydantic import BaseModel,Field


import asyncio

#from selenium import webdriver
#from selenium.webdriver.chrome.service import Service
#from selenium.webdriver.common.by import By
#from selenium.common.exceptions import WebDriverException

import requests
from bs4 import BeautifulSoup
import urllib3.exceptions

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO


import time


###############################################################################################################
#All Models Which Were Used in this programs

#embedding model

embedding_model_name = "sentence-transformers/all-MiniLM-L6-v2"
#embedding_model_name ="BAAI/bge-base-en-v1.5"

model_sentence = SentenceTransformer(embedding_model_name, device="cpu" )

#compression and reranker model of the embedding
cross_encoder_model = HuggingFaceCrossEncoder(model_name="BAAI/bge-reranker-base",model_kwargs={"device": "cpu"})
compressor = CrossEncoderReranker(model=cross_encoder_model, top_n=8)

#LLM MOdel "GEMINI"
API_KEY=os.getenv("GEMINI_API_KEY")

genai.configure(api_key=API_KEY)
model= genai.GenerativeModel("gemini-2.0-flash")

#################################################################################################################
class EmbeddingModel(Embeddings):
    def __init__(self, model_sentence):
        self.model = model_sentence
        
    def embed_documents(self, texts):   
        return [self.model.encode(text) for text in texts]
    
    def embed_query(self, text):
        return self.model.encode(text)
    

################################################################################################

st.title("Automatic Blog Maker")

query= st.text_input("type your blog name")


################################################################################################
#rss_url1 = "https://news.google.com/rss/search?q=Recently+cybersecurity+breaches+OR+cyber+attacks&hl=en-US&gl=US&ceid=US:en"
rss_url="""https://news.google.com/rss/search?q="data+breach"+OR+"cyber+attack"+OR+ransomware+OR+"security+incident"&hl=en-US&gl=US&ceid=US:en"""

feed = feedparser.parse(rss_url)
titles = [entry.title for entry in feed.entries]





#code for scrapping text from all 5 links from google search
###################################################################


urls=[]
def extract_text_url(urls):
    print(f"Searching For {query} on WEB.....") 

    content = []
    for url in urls:
        try:
            response = requests.get(url, timeout=10, headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
            })
            time.sleep(4) 
            
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Extract visible body text
                body = soup.body.get_text(separator=' ', strip=True) if soup.body else ''
                content.append(body)
            else:
                print(f"Failed to retrieve {url} with status code: {response.status_code}")

        except urllib3.exceptions.SSLError as e:
            print(f"SSL error occurred:{url}:- {e}")
        except requests.exceptions.RequestException as e:
            print(f"Request error:{url}:- {e}")
        except Exception as e:
            print(f"General error:{url}:- {e}")

    text = ''.join(content)
    cleaned_text = re.sub(r'\s+', ' ', text).strip()
    return cleaned_text

#################################################################################################################################################################

def chunk_text(cleaned_text):
    text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=800,
    chunk_overlap=100,
    length_function=len,
    is_separator_regex=False
    )
    chunked_text = text_splitter.split_text(cleaned_text)
    final_chunks_documents = text_splitter.create_documents(chunked_text)

    return final_chunks_documents




def vectorstore_retrieval(query, final_chunked_documents):
    try:
        embedding_model= EmbeddingModel(model_sentence)
        vector_store= FAISS.from_documents(final_chunked_documents, embedding_model)
        retriever= vector_store.as_retriever(search_type= "similarity", search_kwargs={"k": 20})
        reranked_text=""
        compression_retriever= ContextualCompressionRetriever(base_compressor=compressor,base_retriever=retriever)
        compressed_document= compression_retriever.invoke(f"Title: , Breif: , Attack Description: , Impact on Systems & Users: ,Technical Description: , Impact Analysis: and  Remediation: for {query}")
        #compressed_document= compression_retriever.get_relevant_documents(f"Details about: {query} ")
        for doc in compressed_document:
            reranked_text = reranked_text + doc.page_content
        return reranked_text
    

    except ImportError as e:
        st.error(f"❌ FAISS or dependency not available: {e}")
        raise

    except Exception as e:
        st.error(f"❌ Error in vectorstore_retrieval(): {e}")
        raise


def generate_knowledge_base(query):
    try:
        cleaned_text= extract_text_url(urls)
        final_chunked_documents= chunk_text(cleaned_text)
        reranked_text= vectorstore_retrieval(query, final_chunked_documents)
        print("done generating knowledge_base")
        return reranked_text

    except Exception as e:
        st.error(f"❌ Error in generate_knowledge_base(): {e}")
        raise
    





def get_llm_response(query,reranked_text):
    try:
        prompt= f"""
        Write a 5-6 page detailed blog post using the following source material:
        {query}
        
        Include Sections With Same Format:
        ## Title: ##
        
        ### Brief Summary: ###
        body text
        
        ### Attack Description: ###
        body text
        
        ### Impact on Systems & Users: ###
        body text
        
        ### Technical Description: ###
        body text
        
        ### Impact Analysis: ###
        body text
        
        ### Remediation Steps(in bulletin points): ###
        body text
        
        Source Material:
        {reranked_text}
        """
    
        output= model.generate_content(prompt)
        return output.text

    except Exception as e:
        st.error("❌ Error in get_llm_response():", e)
        raise




def create_formatted_doc_from_markdown(text_output):
    doc = Document()

    # Define custom styles
    styles = doc.styles

    # Title (##)
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(25)
    title_style.font.bold = True

    # Main Heading (###)
    main_heading_style = styles.add_style('CustomMainHeading', WD_STYLE_TYPE.PARAGRAPH)
    main_heading_style.font.size = Pt(18)
    main_heading_style.font.bold = True

    # Subheading (** text **)
    subheading_style = styles.add_style('CustomSubHeading', WD_STYLE_TYPE.PARAGRAPH)
    subheading_style.font.size = Pt(16)
    subheading_style.font.italic = True

    # Body text
    body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.size = Pt(14)

    lines = text_output.splitlines()
    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Title
        if stripped.startswith("## "):
            doc.add_paragraph(stripped[3:].strip(), style='CustomTitle')

        # Main Heading
        elif stripped.startswith("### "):
            doc.add_paragraph("")
            doc.add_paragraph(stripped[4:].strip(), style='CustomMainHeading')

        # Subheading (** ... ** full line)
        elif re.match(r'^\*\* .+ \*\*$', stripped):
            clean_text = stripped[3:-3].strip()
            doc.add_paragraph(clean_text, style='CustomSubHeading')

        # Bullet point with optional bold beginning
        elif stripped.startswith("* "):
            bullet_line = stripped[2:].strip()
            para = doc.add_paragraph(style='List Bullet')

            if bullet_line.startswith("**"):
                # Match pattern like: **Bold Start:** rest of line
                match = re.match(r'^\*\*(.+?)\*\*(.*)', bullet_line)
                if match:
                    bold_text = match.group(1).strip()
                    rest_text = match.group(2).strip()

                    run = para.add_run(bold_text)
                    run.bold = True
                    run.font.size = Pt(14)

                    if rest_text:
                        run2 = para.add_run(" " + rest_text)
                        run2.font.size = Pt(14)
                else:
                    # Just fallback as body text if not match
                    run = para.add_run(bullet_line)
                    run.font.size = Pt(14)
            else:
                run = para.add_run(bullet_line)
                run.font.size = Pt(14)

        # Body text
        else:
            doc.add_paragraph(stripped, style='CustomBody')

    buffer= BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

    print(f"✅ Document saved as:.........")






def main(query):
    try:
        # searching all the links from google
        url= search(query)
        for i in url:
            urls.append(i)
    
        reranked_text= generate_knowledge_base(query)
        print(reranked_text)
        text_output= get_llm_response(query,reranked_text)

        docx_buffer=create_formatted_doc_from_markdown(text_output)
        return text_output, docx_buffer
        
    except Exception as e:
        st.error(f"❌ Error in main(): {e}")
        raise


def blocking_main(query):
    try:
        return main(query)

    except Exception as e:
        st.error(f"❌ Error in blocking_main(): {e}")
        raise





if st.button(f"Generate Blog.. ✍️"):

    if not query.strip():
        st.warning("please enter some input")
    else:
        with st.spinner("🕒 Generating your blog... please wait......"):
            try:
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(blocking_main, query)
                    text_output, docx_buffer = future.result()
    
                st.success("Blog Generated Successfully \n You can download it from Below")
    
                #creating a good file name for word document#
                modified_filename=re.sub(r'[^a-zA-Z0-9\s]', '_', query)
                modified_filename= "_".join(modified_filename.split())
    
    
                st.download_button(
                    label="📥 Download Blog as Word Document",
                    data=docx_buffer,
                    file_name=f"{modified_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            
            except Exception as e:
                st.error("❌ Something went wrong during blog generation.")
                st.exception(e)


st.subheader(f"Recent Cyber Attacks and Breaches 🛡️")
for i, title in enumerate(titles[:30],1):
    st.write(f"🔴({i})--> {title}")
            









